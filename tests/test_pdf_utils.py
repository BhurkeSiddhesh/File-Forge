from pathlib import Path
import pytest
import pikepdf
import fitz
from docx import Document
from reportlab.pdfgen import canvas
from scripts.pdf_utils import (
    remove_pdf_password, pdf_to_docx, extract_pdf_pages, compress_pdf, extract_pdf_text,
    _parse_page_selection, _inspect_text_layer, pdf_to_word_ai,
)
import scripts.pdf_utils as pdf_utils_module
import scripts.ocr_engine as ocr_engine

def test_remove_pdf_password(locked_pdf, tmp_path):
    """Test removing password from a PDF."""
    input_path = locked_pdf["path"]
    password = locked_pdf["password"]

    # Use tmp_path as the output directory
    output_path_str = remove_pdf_password(str(input_path), password, str(tmp_path))
    output_path = Path(output_path_str)

    assert output_path.exists()
    assert output_path.name == "locked_forgefiles.org.pdf"

    # Verify the file can be opened without password
    with pikepdf.open(output_path) as pdf:
        assert len(pdf.pages) > 0

def test_pdf_to_docx(sample_pdf, tmp_path):
    """Test converting PDF to DOCX."""
    output_path_str = pdf_to_docx(str(sample_pdf), str(tmp_path))
    output_path = Path(output_path_str)

    assert output_path.exists()
    assert output_path.suffix == ".docx"
    assert output_path.stem == f"{sample_pdf.stem}_forgefiles.org"

def test_extract_pdf_pages(multi_page_pdf, tmp_path):
    """Test extracting selected pages from a PDF."""
    output_path_str = extract_pdf_pages(str(multi_page_pdf), str(tmp_path), "1,3-4")
    output_path = Path(output_path_str)

    assert output_path.exists()
    assert output_path.name == "multi_sample_forgefiles.org.pdf"

    with pikepdf.open(output_path) as pdf:
        assert len(pdf.pages) == 3


def test_extract_pdf_text_creates_txt_file(sample_pdf, tmp_path):
    """Text extraction writes embedded PDF text to a TXT file."""
    output_path_str = extract_pdf_text(str(sample_pdf), str(tmp_path), use_ocr=False)
    output_path = Path(output_path_str)

    assert output_path.exists()
    assert output_path.name == "sample_forgefiles.org.txt"
    assert "Hello, this is a test PDF." in output_path.read_text(encoding="utf-8")


def test_extract_pdf_text_with_password(locked_pdf, tmp_path):
    """Text extraction supports password-protected PDFs."""
    output_path_str = extract_pdf_text(
        str(locked_pdf["path"]),
        str(tmp_path),
        password=locked_pdf["password"],
        use_ocr=False,
    )

    assert "Hello, this is a test PDF." in Path(output_path_str).read_text(encoding="utf-8")


def test_extract_pdf_text_encrypted_no_password_raises(locked_pdf, tmp_path):
    """Text extraction without password for encrypted PDF raises ValueError."""
    with pytest.raises(ValueError, match="password"):
        extract_pdf_text(str(locked_pdf["path"]), str(tmp_path), use_ocr=False)


def test_extract_pdf_text_blank_pdf_without_ocr_raises(tmp_path):
    """A PDF without embedded text raises a user-friendly error when OCR is disabled."""
    from reportlab.pdfgen import canvas

    blank_pdf = tmp_path / "blank.pdf"
    c = canvas.Canvas(str(blank_pdf))
    c.showPage()
    c.save()

    with pytest.raises(ValueError, match="No text"):
        extract_pdf_text(str(blank_pdf), str(tmp_path), use_ocr=False)


def test_remove_pdf_password_wrong_password(locked_pdf, tmp_path):
    """Wrong password on an open-password PDF raises an error (no false unlock)."""
    with pytest.raises(Exception):
        remove_pdf_password(str(locked_pdf["path"]), "wrong_password", str(tmp_path))


def _make_owner_restricted_pdf(sample_pdf, dest):
    """An owner-restricted PDF: permission restrictions but NO open/user password,
    so it opens with an empty password (the common 'Unlock PDF' case)."""
    with pikepdf.open(sample_pdf) as pdf:
        pdf.save(
            dest,
            encryption=pikepdf.Encryption(
                user="",  # ggignore  — no open password
                owner="owner-secret",  # ggignore
                allow=pikepdf.Permissions(extract=False, modify_other=False, print_highres=False),
            ),
        )
    return dest


def _pdf_text(path):
    doc = fitz.open(str(path))
    try:
        return "".join(page.get_text() for page in doc)
    finally:
        doc.close()


def test_remove_pdf_password_owner_restricted_with_wrong_password(sample_pdf, tmp_path):
    """Owner-restricted PDFs (no open password) unlock even when the caller passes
    a wrong/irrelevant password, matching mainstream Unlock PDF tools. Content must
    be preserved and the output must be fully decrypted."""
    src = _make_owner_restricted_pdf(sample_pdf, tmp_path / "restricted.pdf")
    baseline_text = _pdf_text(sample_pdf)

    out = Path(remove_pdf_password(str(src), "a-guess-that-is-wrong", str(tmp_path)))

    assert out.exists()
    with pikepdf.open(out) as pdf:  # opens with no password → truly unlocked
        assert not pdf.is_encrypted
    # Extracted content is byte/hash-identical to the source (zero fidelity loss).
    assert _pdf_text(out) == baseline_text


def test_remove_pdf_password_owner_restricted_removes_restrictions(sample_pdf, tmp_path):
    """After unlocking, the owner-level permission restrictions are gone."""
    src = _make_owner_restricted_pdf(sample_pdf, tmp_path / "restricted2.pdf")

    out = Path(remove_pdf_password(str(src), "", str(tmp_path)))

    with pikepdf.open(out) as pdf:
        assert not pdf.is_encrypted
        # A non-encrypted PDF imposes no permission restrictions.
        assert pdf.allow.extract and pdf.allow.print_highres


def test_remove_pdf_password_owner_restricted_with_owner_password(sample_pdf, tmp_path):
    """Supplying the correct owner password also unlocks (first-try path)."""
    src = _make_owner_restricted_pdf(sample_pdf, tmp_path / "restricted3.pdf")

    out = Path(remove_pdf_password(str(src), "owner-secret", str(tmp_path)))  # ggignore

    with pikepdf.open(out) as pdf:
        assert not pdf.is_encrypted


def test_pdf_to_docx_with_password(locked_pdf, tmp_path):
    """Converting a password-protected PDF works when the correct password is supplied."""
    output_path_str = pdf_to_docx(str(locked_pdf["path"]), str(tmp_path), password=locked_pdf["password"])
    output_path = Path(output_path_str)
    assert output_path.exists()
    assert output_path.suffix == ".docx"


def test_pdf_to_docx_encrypted_no_password(locked_pdf, tmp_path):
    """Converting an encrypted PDF without a password raises ValueError."""
    with pytest.raises(ValueError, match="password"):
        pdf_to_docx(str(locked_pdf["path"]), str(tmp_path))


# ---------------------------------------------------------------------------
# _parse_page_selection tests
# ---------------------------------------------------------------------------

def test_parse_page_selection_single_page():
    """Single page number returns correct zero-based index."""
    assert _parse_page_selection("1", 5) == [0]


def test_parse_page_selection_multiple_pages():
    """Comma-separated pages return correct zero-based indices in order."""
    assert _parse_page_selection("1,3,5", 5) == [0, 2, 4]


def test_parse_page_selection_range():
    """Hyphen range returns all pages in range."""
    assert _parse_page_selection("2-4", 5) == [1, 2, 3]


def test_parse_page_selection_mixed():
    """Mixed single pages and ranges return correct indices."""
    assert _parse_page_selection("1,3-4", 4) == [0, 2, 3]


def test_parse_page_selection_all():
    """'all' keyword returns every page index."""
    assert _parse_page_selection("all", 4) == [0, 1, 2, 3]


def test_parse_page_selection_deduplication():
    """Duplicate page numbers appear only once in the result."""
    result = _parse_page_selection("1,1,2", 5)
    assert result == [0, 1]


def test_parse_page_selection_none_raises():
    """None input raises ValueError."""
    with pytest.raises(ValueError):
        _parse_page_selection(None, 5)


def test_parse_page_selection_empty_string_raises():
    """Empty string raises ValueError."""
    with pytest.raises(ValueError):
        _parse_page_selection("", 5)


def test_parse_page_selection_out_of_bounds_raises():
    """Page number beyond document length raises ValueError."""
    with pytest.raises(ValueError, match="exceeds"):
        _parse_page_selection("10", 5)


def test_parse_page_selection_zero_page_raises():
    """Page number 0 is invalid (pages are 1-based)."""
    with pytest.raises(ValueError):
        _parse_page_selection("0", 5)


def test_parse_page_selection_inverted_range_raises():
    """Range where start > end raises ValueError."""
    with pytest.raises(ValueError):
        _parse_page_selection("5-2", 5)


def test_parse_page_selection_non_numeric_raises():
    """Non-numeric page token raises ValueError."""
    with pytest.raises(ValueError):
        _parse_page_selection("abc", 5)


def test_parse_page_selection_whitespace_trimmed():
    """Leading/trailing whitespace around tokens is handled gracefully."""
    assert _parse_page_selection(" 1 , 3 ", 5) == [0, 2]


# ---------------------------------------------------------------------------
# compress_pdf tests
# ---------------------------------------------------------------------------

def test_compress_pdf_returns_dict_keys(sample_pdf, tmp_path):
    """compress_pdf returns a dict with the expected keys."""
    result = compress_pdf(str(sample_pdf), str(tmp_path), level='low')
    assert "output_path" in result
    assert "original_size" in result
    assert "compressed_size" in result
    assert "reduction_pct" in result


def test_compress_pdf_output_file_exists(sample_pdf, tmp_path):
    """compress_pdf creates an output file."""
    result = compress_pdf(str(sample_pdf), str(tmp_path), level='low')
    assert Path(result["output_path"]).exists()


def test_compress_pdf_output_is_valid_pdf(sample_pdf, tmp_path):
    """Output of compress_pdf is a valid, openable PDF."""
    result = compress_pdf(str(sample_pdf), str(tmp_path), level='low')
    with pikepdf.open(result["output_path"]) as pdf:
        assert len(pdf.pages) > 0


def test_compress_pdf_medium_level(sample_pdf, tmp_path):
    """compress_pdf works with level='medium'."""
    result = compress_pdf(str(sample_pdf), str(tmp_path), level='medium')
    assert Path(result["output_path"]).exists()


def test_compress_pdf_high_level(sample_pdf, tmp_path):
    """compress_pdf works with level='high'."""
    result = compress_pdf(str(sample_pdf), str(tmp_path), level='high')
    assert Path(result["output_path"]).exists()


def test_compress_pdf_with_password(locked_pdf, tmp_path):
    """compress_pdf decrypts and compresses a password-protected PDF."""
    result = compress_pdf(str(locked_pdf["path"]), str(tmp_path), level='low', password=locked_pdf["password"])
    assert Path(result["output_path"]).exists()


def test_compress_pdf_encrypted_no_password_raises(locked_pdf, tmp_path):
    """compress_pdf without password for encrypted PDF raises ValueError."""
    with pytest.raises(ValueError, match="password"):
        compress_pdf(str(locked_pdf["path"]), str(tmp_path))


# ---------------------------------------------------------------------------
# pdf_to_word_ai: text-layer detection + routing
#
# Regression coverage for a real bug: pdf_to_word_ai used to rasterize and
# OCR every page unconditionally, even PDFs with a perfectly clean embedded
# text layer, which on ARM deployments (OCR_BACKEND=rapidocr, no layout
# support) produced mangled output (dropped spaces, "AI"->"Al", reordered
# bullets) that was strictly worse than the standard converter. It should
# never touch OCR when a usable text layer already exists.
# ---------------------------------------------------------------------------

def test_inspect_text_layer_classifies_text_and_scanned_pages(text_rich_pdf, scanned_like_pdf):
    """A text-native PDF is classified as having a usable text layer; an
    image-only PDF is not."""
    doc = fitz.open(str(text_rich_pdf))
    try:
        report = _inspect_text_layer(doc)
    finally:
        doc.close()
    assert report["has_usable_text_layer"] is True
    assert report["fraction_with_text"] == 1.0

    doc = fitz.open(str(scanned_like_pdf))
    try:
        report = _inspect_text_layer(doc)
    finally:
        doc.close()
    assert report["has_usable_text_layer"] is False
    assert report["fraction_with_text"] == 0.0


def test_pdf_to_word_ai_skips_ocr_for_text_pdf(text_rich_pdf, tmp_path, monkeypatch):
    """pdf_to_word_ai must not rasterize+OCR a PDF that already has a usable
    text layer, regardless of which OCR backend is configured."""
    def _fail(*args, **kwargs):
        raise AssertionError("OCR should not run for a text-based PDF")

    monkeypatch.setattr(pdf_utils_module, "_pdf_to_word_ocr_fallback", _fail)
    monkeypatch.setattr(pdf_utils_module, "_pdf_to_word_paddle_impl", _fail)
    monkeypatch.setattr(pdf_utils_module, "_pdf_to_word_hybrid_impl", _fail)

    methods = []
    output_path = pdf_to_word_ai(
        str(text_rich_pdf), str(tmp_path), method_callback=methods.append,
    )
    assert Path(output_path).exists()
    assert methods == ["text_layer"]


def test_pdf_to_word_ai_text_pdf_succeeds_even_with_ai_disabled(text_rich_pdf, tmp_path, monkeypatch):
    """A text-based PDF converts fine even when AI/OCR is disabled server-wide -
    it never needed OCR in the first place."""
    monkeypatch.setenv("DISABLE_AI", "1")
    ocr_engine.reset_engine()

    output_path = pdf_to_word_ai(str(text_rich_pdf), str(tmp_path))
    assert Path(output_path).exists()


def test_pdf_to_word_ai_scanned_pdf_still_errors_when_ai_disabled(scanned_like_pdf, tmp_path, monkeypatch):
    """A genuinely scanned/image-only PDF still requires AI/OCR to be enabled."""
    monkeypatch.setenv("DISABLE_AI", "1")
    ocr_engine.reset_engine()

    with pytest.raises(ValueError):
        pdf_to_word_ai(str(scanned_like_pdf), str(tmp_path))


def test_pdf_to_word_ai_output_has_correct_spacing_and_no_corrupted_ai(text_rich_pdf, tmp_path):
    """The routed (standard, text-layer-based) conversion preserves word
    spacing and doesn't corrupt 'AI' - the exact symptoms of the OCR bug."""
    output_path = pdf_to_word_ai(str(text_rich_pdf), str(tmp_path))
    doc = Document(output_path)
    text = "\n".join(p.text for p in doc.paragraphs)

    assert "Senior Analytics Consultant" in text
    assert "AI" in text
    assert "Al " not in text and "Al." not in text


def test_pdf_to_word_ai_hybrid_routes_ocr_only_to_scanned_pages(tmp_path, monkeypatch):
    """A mixed PDF (one text page, one scanned page) should extract the text
    page natively and OCR only the scanned page."""
    from PIL import Image
    from unittest.mock import MagicMock

    d = tmp_path / "mixed_src"
    d.mkdir()
    file_path = d / "mixed.pdf"

    c = canvas.Canvas(str(file_path))
    c.drawString(72, 750, "This page already has plenty of real embedded text content.")
    c.showPage()
    img_path = d / "page.png"
    Image.new("RGB", (200, 200), color="white").save(img_path)
    c.drawImage(str(img_path), 0, 0, width=200, height=200)
    c.showPage()
    c.save()

    fake_engine = MagicMock()
    fake_engine.name = "fake"
    fake_engine.supports_layout = False
    fake_engine.recognize.return_value = [
        {"text": "Scanned line", "bbox": [[0, 0], [100, 0], [100, 20], [0, 20]]}
    ]

    monkeypatch.setattr(ocr_engine, "get_ocr_engine", lambda *a, **k: fake_engine)

    methods = []
    output_dir = tmp_path / "out"
    output_dir.mkdir()
    output_path = pdf_to_word_ai(str(file_path), str(output_dir), method_callback=methods.append)

    assert Path(output_path).exists()
    assert methods == ["ocr_hybrid"]
    # Only the scanned page (page 2) should have gone through OCR.
    assert fake_engine.recognize.call_count == 1

    doc = Document(output_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "This page already has plenty of real embedded text content." in text
    assert "Scanned line" in text
